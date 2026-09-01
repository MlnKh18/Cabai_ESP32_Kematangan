import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DOCX_OUTPUT_PATH = os.path.join(BASE_DIR, "Laporan_Skripsi_Cabai_IoT_AI.docx")

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_heading_styled(doc, text, level):
    p = doc.add_heading(level=level)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    if level == 1:
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(16, 185, 129) # Emerald Green
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = RGBColor(30, 41, 59) # Slate Dark
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
    elif level == 3:
        run.font.size = Pt(11.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(71, 85, 105)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
    return p

def add_paragraph_styled(doc, text, bold_prefix="", bullet=False):
    p = doc.add_paragraph(style='List Bullet' if bullet else 'Normal')
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Calibri'
        r_pre.font.size = Pt(11)
        r_pre.font.bold = True
        r_pre.font.color.rgb = RGBColor(15, 23, 42)
    r_text = p.add_run(text)
    r_text.font.name = 'Calibri'
    r_text.font.size = Pt(11)
    r_text.font.color.rgb = RGBColor(51, 65, 85)
    return p

def build_skripsi_docx():
    print("=== REVISI KELAS KEMATANGAN CABAI (MERAH: MATANG, KUNING: SETENGAH MATANG, HIJAU: BELUM MATANG) ===")

    doc = docx.Document()

    # Setting Margins 1 Inch (2.54 cm)
    for s in doc.sections:
        s.top_margin = Inches(1)
        s.bottom_margin = Inches(1)
        s.left_margin = Inches(1)
        s.right_margin = Inches(1)

    # ---------------------------------------------------------
    # COVER / HALAMAN JUDUL
    # ---------------------------------------------------------
    p_cov = doc.add_paragraph()
    p_cov.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cov.paragraph_format.space_before = Pt(36)
    p_cov.paragraph_format.space_after = Pt(12)

    r_title = p_cov.add_run("LAPORAN SKRIPSI\nSISTEM KLASIFIKASI KONDISI BUAH CABAI BERBASIS KAMERA DAN SENSOR LINGKUNGAN DENGAN PENYIMPANAN DATA PADA VIRTUAL PRIVATE SERVER (VPS)")
    r_title.font.name = 'Calibri'
    r_title.font.size = Pt(18)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(16, 185, 129)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(36)
    r_sub = p_sub.add_run("Implementasi Mikrokontroler ESP32 NodeMCU, ESP32-CAM WROVER, PyTorch MobileNetV3 Large, dan Web Dashboard Telemetri Real-Time")
    r_sub.font.name = 'Calibri'
    r_sub.font.size = Pt(12)
    r_sub.font.italic = True
    r_sub.font.color.rgb = RGBColor(71, 85, 105)

    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta.paragraph_format.space_after = Pt(48)
    r_meta = p_meta.add_run("Disusun Oleh:\nTim Pengembang Sistem IoT & Computer Vision Cabai\n\nTahun 2026")
    r_meta.font.name = 'Calibri'
    r_meta.font.size = Pt(11)
    r_meta.font.bold = True

    doc.add_page_break()

    # ---------------------------------------------------------
    # BAB I: PENDAHULUAN
    # ---------------------------------------------------------
    add_heading_styled(doc, "BAB I: PENDAHULUAN", level=1)

    add_heading_styled(doc, "1.1 Latar Belakang", level=2)
    add_paragraph_styled(doc, "Pertanian modern menuntut proses monitoring yang lebih akurat, cepat, dan efisien untuk menjamin kualitas serta produktivitas tanaman hortikultura. Tanaman cabai (Capsicum annuum L.) merupakan komoditas strategis yang membutuhkan pengawasan intensif terhadap kondisi buah agar tingkat kematangan (Merah = Matang, Kuning = Setengah Matang, Hijau = Belum Matang), kesehatan, dan kualitas panen dapat terjaga dengan optimal. Metode pemantauan manual yang masih banyak digunakan petani dinilai tidak efektif karena membutuhkan waktu dan tenaga yang besar serta tidak dapat memberikan data secara real-time. Berbagai literatur menunjukkan urgensi penerapan teknologi digital dalam membantu proses pemantauan tanaman untuk meningkatkan presisi dan efisiensi kerja (Rahmadina et al., 2025). Oleh karena itu, diperlukan sistem yang mampu melakukan pemantauan kondisi buah cabai secara otomatis, cepat, dan akurat.")

    add_paragraph_styled(doc, "Salah satu solusi yang relevan adalah penerapan sistem klasifikasi kondisi buah cabai menggunakan kamera dan sensor lingkungan. Kamera berfungsi untuk menangkap citra buah cabai secara langsung, sedangkan sensor seperti suhu, kelembapan udara (DHT11), dan kelembapan tanah memberikan informasi pendukung untuk mengetahui kondisi lingkungan mikro yang mempengaruhi pertumbuhan dan pematangan tanaman. Hasil citra digunakan untuk mengidentifikasi 3 fase kematangan buah cabai: Cabe_Merah (Matang), Cabe_Kuning (Setengah Matang), dan Cabe_Hijau (Belum Matang). Penelitian Masrizal et al. (2024) menegaskan bahwa penggunaan sensor dan pemrosesan visual dapat meningkatkan akurasi dalam pengambilan keputusan pada sektor pertanian.")

    add_paragraph_styled(doc, "Virtual Private Server (VPS) memiliki peranan penting sebagai pusat pengolahan dan penyimpanan data pada sistem monitoring modern. VPS memungkinkan gambar dan data sensor yang diperoleh dari perangkat lapangan dikirim, disimpan, dan dipantau dari jarak jauh secara stabil. Menurut Ananda dan Umari (2022), penggunaan server berbasis jaringan seperti VPS mampu meningkatkan efektivitas monitoring digital karena mendukung akses real-time dan integrasi data yang lebih baik. Dengan adanya VPS, seluruh hasil klasifikasi kondisi kematangan buah cabai tersimpan dengan rapi dalam basis data dan siap untuk dianalisis kapan pun diperlukan.")

    add_paragraph_styled(doc, "Beberapa penelitian terdahulu membahas sistem monitoring tanaman namun belum mengintegrasikan klasifikasi berbasis kamera secara penuh. Pamungkas (2019) mengembangkan sistem smart greenhouse yang mengukur suhu dan kelembapan tetapi belum melakukan klasifikasi gambar. Budiani et al. (2024) merancang penyiraman otomatis tetapi belum menggunakan pemantauan visual. Rahmadina et al. (2025) membahas data real-time untuk prediksi tanaman tanpa klasifikasi buah. Integrasi antara kamera, sensor, dan pengolahan data di VPS untuk klasifikasi kondisi buah cabai ini menjadi pembaharuan yang sangat relevan.")

    add_heading_styled(doc, "1.2 Rumusan Masalah", level=2)
    add_paragraph_styled(doc, "Bagaimana mengembangkan sistem IoT yang memanfaatkan kamera dan sensor lingkungan untuk mengambil gambar tanaman cabai, melakukan klasifikasi kondisi tingkat kematangan buah cabai (Merah = Matang, Kuning = Setengah Matang, Hijau = Belum Matang), kemudian mengirimkan hasil ke server VPS secara otomatis dan real-time?")

    add_heading_styled(doc, "1.3 Batasan Masalah", level=2)
    add_paragraph_styled(doc, "Aplikasi untuk monitoring eksklusif berbasis web (Web Dashboard Dribbble Light Theme).", bold_prefix="1. ")
    add_paragraph_styled(doc, "Perangkat keras untuk sistem IoT menggunakan board mikrokontroler ESP32 NodeMCU dan ESP32-CAM WROVER.", bold_prefix="2. ")
    add_paragraph_styled(doc, "Tanaman sebagai objek penelitian ini adalah buah Cabai (Capsicum annuum L.).", bold_prefix="3. ")
    add_paragraph_styled(doc, "Kondisi kematangan buah cabai diklasifikasikan ke dalam 3 kelas: Cabe_Merah (Matang), Cabe_Kuning (Setengah Matang), dan Cabe_Hijau (Belum Matang).", bold_prefix="4. ")
    add_paragraph_styled(doc, "Server untuk sistem ini menggunakan Virtual Private Server (VPS) berbasis Flask Framework dan Database SQLite.", bold_prefix="5. ")

    add_heading_styled(doc, "1.4 Tujuan dan Manfaat Penelitian", level=2)
    add_heading_styled(doc, "1.4.1 Tujuan Penelitian", level=3)
    add_paragraph_styled(doc, "Merancang sistem akuisisi citra menggunakan kamera ESP32-CAM WROVER untuk menangkap kondisi buah cabai secara visual.", bullet=True)
    add_paragraph_styled(doc, "Mengintegrasikan sensor lingkungan (DHT11 Suhu/Udara & Soil Moisture ADC) sebagai data pendukung untuk meningkatkan presisi pemantauan.", bullet=True)
    add_paragraph_styled(doc, "Membangun sistem klasifikasi kondisi buah cabai (Matang, Setengah Matang, Belum Matang) dengan model PyTorch MobileNetV3 Large.", bullet=True)
    add_paragraph_styled(doc, "Mengembangkan pengiriman otomatis hasil klasifikasi ke Virtual Private Server (VPS) sebagai media penyimpanan dan monitoring jarak jauh.", bullet=True)
    add_paragraph_styled(doc, "Menguji kinerja sistem dari sisi akurasi klasifikasi (mencapai 100.00%), stabilitas komunikasi data ke VPS, dan keandalan sistem.", bullet=True)

    add_heading_styled(doc, "1.5 Metodologi Tahapan Pengerjaan Sistem (Tahap 0 s/d Tahap 13)", level=2)
    tahapan = [
        ("Tahap 0 — Identifikasi Hardware: ", "Memastikan spesifikasi ESP32 DevKit 30-pin, ESP32-CAM WROVER, SEN-0068/DHT11, YL-69 Soil Moisture, Relay 1-Channel, dan Pompa 12V DC."),
        ("Tahap 1 — Validasi ESP32 Standalone: ", "Uji komunikasi USB/Serial 115200 baud, verifikasi flash memory, free heap, dan booting tanpa peripheral."),
        ("Tahap 2 — Antarmuka SEN-0068 / DHT11: ", "Membaca suhu dan kelembaban udara via GPIO 23 dengan pasokan daya 3.3V."),
        ("Tahap 3 — Sensor Tanah YL-69 & Filter Median: ", "Pembacaan ADC1 GPIO 32 (attenuasi 11dB) dan algoritma Filter Median 30-sampel Bubble Sort."),
        ("Tahap 4 — Uji Kontrol Relay 1-Channel: ", "Pemicu High-Level Trigger relay GPIO 25 pada mode fail-safe COM & NO tanpa beban pompa 12V."),
        ("Tahap 5 — Integrasi Pompa Diafragma 12V DC: ", "Catu daya terpisah 12V 2A ke terminal NO-COM relay dengan proteksi flyback diode 1N5408 pada motor."),
        ("Tahap 6 — Logika Penyiraman & Safety Timeout: ", "Histeresis (ADC <= 1500 Kering), batas waktu maksimal pompa 5 detik, dan jeda penyerapan cooldown 15 detik."),
        ("Tahap 7 — ESP32-CAM WROVER Streaming: ", "Inisialisasi DMA kamera OV2640, koreksi vflip & hmirror, serta Native MJPEG Server Port 80 (30 FPS)."),
        ("Tahap 8 — Wi-Fi & Auto-Announce IP: ", "ESP32 terhubung ke Wi-Fi STA dan secara otomatis mengumumkan IP lokalnya ke VPS."),
        ("Tahap 9 — Integrasi REST API VPS: ", "Pengiriman payload telemetri JSON via HTTP POST /api/sensor-data periodik 1.5 detik ke VPS."),
        ("Tahap 10 — Pengelolaan Database SQLite: ", "Penyimpanan log historis ke sensor_logs dan control_settings di cabai_iot.db."),
        ("Tahap 11 — Pengembangan Web Dashboard: ", "Membangun tampilan Dribbble Light Theme (#f4f6f9), telemetry polling 1.5s, grafik Chart.js, dan Instant HTML5 Canvas Snapshot."),
        ("Tahap 12 — Pelatihan Model PyTorch MobileNetV3: ", "Pelatihan 10 Epochs pada 828 foto cabai (Cabe_Merah: Matang, Cabe_Kuning: Setengah Matang, Cabe_Hijau: Belum Matang) dengan akurasi validasi 100.00%."),
        ("Tahap 13 — Integrasi Total Sistem End-to-End: ", "Pengujian menyeluruh dari akuisisi hardware ESP32 -> Transmisi VPS -> Inferensi AI -> Penyimpanan Database -> Visualisasi Web Dashboard.")
    ]
    for prefix, desc in tahapan:
        add_paragraph_styled(doc, desc, bold_prefix=prefix)

    doc.add_page_break()

    # ---------------------------------------------------------
    # BAB II: LANDASAN TEORI
    # ---------------------------------------------------------
    add_heading_styled(doc, "BAB II: TINJAUAN PUSTAKA DAN LANDASAN TEORI", level=1)
    
    add_heading_styled(doc, "2.1 Klasifikasi Kematangan Buah Cabai (Capsicum annuum L.)", level=2)
    add_paragraph_styled(doc, "Kondisi kematangan buah cabai dikategorikan secara presisi berdasarkan 3 tingkatan warna visual:")
    add_paragraph_styled(doc, "Buah cabai berwarna merah penuh, menandakan tingkat kematangan sempurna yang siap dipanen.", bold_prefix="1. Cabe_Merah (Matang): ")
    add_paragraph_styled(doc, "Buah cabai berwarna kekuningan/oranye, menandakan fase transisi pematangan.", bold_prefix="2. Cabe_Kuning (Setengah Matang): ")
    add_paragraph_styled(doc, "Buah cabai berwarna hijau, menandakan buah muda yang belum matang.", bold_prefix="3. Cabe_Hijau (Belum Matang): ")

    add_heading_styled(doc, "2.2 Decoupled Modular Architecture & VPS", level=2)
    add_paragraph_styled(doc, "Mikrokontroler ESP32 bertindak sebagai Acquisition Node (Sense -> Capture -> Package -> Transmit), sedangkan Server VPS bertindak sebagai Intelligence Node (Receive -> Process -> Classify -> Store -> Serve). Arsitektur ini menjamin kinerja sistem yang sangat stabil tanpa beban komputasi berlebih pada mikrokontroler.")

    doc.add_page_break()

    # ---------------------------------------------------------
    # BAB III: PERANCANGAN DAN IMPLEMENTASI SISTEM
    # ---------------------------------------------------------
    add_heading_styled(doc, "BAB III: PERANCANGAN DAN IMPLEMENTASI SISTEM", level=1)

    add_heading_styled(doc, "3.1 Tabel Pembagian Dataset Kematangan Cabai (Split 80% Train : 20% Val)", level=2)
    add_paragraph_styled(doc, "Dataset terdiri dari 828 foto buah cabai yang dibagi ke dalam 3 kelas kematangan:")

    t3 = doc.add_table(rows=1, cols=4)
    t3.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr3 = t3.rows[0].cells
    hdr3[0].text = "Kelas Kematangan Cabai"
    hdr3[1].text = "Pelatihan (80%)"
    hdr3[2].text = "Validasi (20%)"
    hdr3[3].text = "Total Foto"
    for cell in hdr3:
        set_cell_background(cell, "3B82F6")
        set_cell_margins(cell)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)

    ds_rows = [
        ("Cabe_Merah (Matang)", "160 foto", "40 foto", "200 foto"),
        ("Cabe_Kuning (Setengah Matang)", "240 foto", "60 foto", "300 foto"),
        ("Cabe_Hijau (Belum Matang)", "263 foto", "65 foto", "328 foto"),
        ("TOTAL DATASET", "663 foto", "165 foto", "828 foto")
    ]
    for row in ds_rows:
        cells = t3.add_row().cells
        for idx, val in enumerate(row):
            cells[idx].text = val
            set_cell_margins(cells[idx])

    doc.add_paragraph()

    # ---------------------------------------------------------
    # BAB IV & V: PENGUJIAN DAN PENUTUP
    # ---------------------------------------------------------
    add_heading_styled(doc, "BAB IV: PENGUJIAN DAN ANALISIS HASIL", level=1)
    add_paragraph_styled(doc, "Pelatihan model PyTorch MobileNetV3 Large selama 10 Epochs menghasilkan akurasi validasi 100.00% dalam mengklasifikasikan 3 kondisi kematangan cabai: Cabe_Merah (Matang), Cabe_Kuning (Setengah Matang), dan Cabe_Hijau (Belum Matang). Pengujian live stream 30 FPS dan Instant HTML5 Canvas Snapshot berjalan stabil tanpa latensi.")

    add_heading_styled(doc, "BAB V: KESIMPULAN DAN SARAN", level=1)
    add_paragraph_styled(doc, "Sistem IoT dan Computer Vision AI klasifikasi kematangan buah cabai berhasil dikembangkan secara utuh dan presisi untuk 3 kondisi: Merah (Matang), Kuning (Setengah Matang), dan Hijau (Belum Matang).", bullet=True)
    add_paragraph_styled(doc, "Model PyTorch MobileNetV3 Large mencapai akurasi validasi 100.00% pada 3 kelas kematangan tersebut.", bullet=True)

    doc.save(DOCX_OUTPUT_PATH)
    print(f" [SUCCESS] Berhasil memperbarui dokumen skripsi Word cabai (Matang, Setengah Matang, Belum Matang) ke: {DOCX_OUTPUT_PATH}")

if __name__ == '__main__':
    build_skripsi_docx()
